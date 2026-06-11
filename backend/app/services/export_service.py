import re
from dataclasses import dataclass
from io import BytesIO
from typing import Literal
from urllib.parse import quote

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, RGBColor
from sqlalchemy.orm import Session

from app.models.resume_version import ResumeVersion
from app.repositories.job_description_repository import JobDescriptionRepository
from app.repositories.resume_version_repository import ResumeVersionRepository

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DocxTemplateName = Literal["standard", "modern", "compact"]
SUPPORTED_DOCX_TEMPLATES: tuple[DocxTemplateName, ...] = ("standard", "modern", "compact")
MAX_EXPORT_FILENAME_STEM_LENGTH = 80
WINDOWS_RESERVED_FILENAMES = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    "COM1",
    "COM2",
    "COM3",
    "COM4",
    "COM5",
    "COM6",
    "COM7",
    "COM8",
    "COM9",
    "LPT1",
    "LPT2",
    "LPT3",
    "LPT4",
    "LPT5",
    "LPT6",
    "LPT7",
    "LPT8",
    "LPT9",
}


@dataclass(frozen=True)
class BinaryExport:
    content: bytes
    filename: str
    content_disposition: str


@dataclass(frozen=True)
class MarkdownBlock:
    kind: Literal["heading1", "heading2", "heading3", "bullet", "paragraph", "spacer"]
    text: str = ""


class ExportNotFoundError(RuntimeError):
    pass


class DocxTemplateRenderer:
    name: DocxTemplateName = "standard"
    normal_size = 10.5
    heading_sizes = (18, 14, 12)
    normal_space_after = 4
    bullet_space_after = 2
    heading_space_before = 8
    heading_space_after = 4
    accent_color: RGBColor | None = None

    def render(self, blocks: list[MarkdownBlock]) -> bytes:
        document = Document()
        self._apply_styles(document)

        for block in blocks:
            self._add_block(document, block)

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _apply_styles(self, document: DocxDocument) -> None:
        styles = document.styles
        normal_style = styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(self.normal_size)
        normal_style.paragraph_format.space_after = Pt(self.normal_space_after)

        for style_name, size in zip(
            ("Heading 1", "Heading 2", "Heading 3"),
            self.heading_sizes,
            strict=True,
        ):
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = True
            if self.accent_color is not None and style_name in {"Heading 1", "Heading 2"}:
                style.font.color.rgb = self.accent_color
            style.paragraph_format.space_before = Pt(self.heading_space_before)
            style.paragraph_format.space_after = Pt(self.heading_space_after)

    def _add_block(self, document: DocxDocument, block: MarkdownBlock) -> None:
        if block.kind == "spacer":
            paragraph = document.add_paragraph()
        elif block.kind == "heading1":
            paragraph = document.add_heading(block.text, level=1)
        elif block.kind == "heading2":
            paragraph = document.add_heading(block.text, level=2)
        elif block.kind == "heading3":
            paragraph = document.add_heading(block.text, level=3)
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(block.text, style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(self.bullet_space_after)
        else:
            paragraph = document.add_paragraph(block.text)
            paragraph.paragraph_format.space_after = Pt(self.normal_space_after)

        self._after_add_block(paragraph, block)

    def _after_add_block(self, paragraph: object, block: MarkdownBlock) -> None:
        return None


class StandardDocxTemplate(DocxTemplateRenderer):
    name = "standard"


class ModernDocxTemplate(DocxTemplateRenderer):
    name = "modern"
    normal_size = 10.5
    heading_sizes = (20, 15, 12)
    normal_space_after = 6
    bullet_space_after = 4
    heading_space_before = 10
    heading_space_after = 6
    accent_color = RGBColor(31, 78, 121)


class CompactDocxTemplate(DocxTemplateRenderer):
    name = "compact"
    normal_size = 9.5
    heading_sizes = (16, 12, 10.5)
    normal_space_after = 1
    bullet_space_after = 0
    heading_space_before = 4
    heading_space_after = 1


DOCX_TEMPLATE_RENDERERS: dict[DocxTemplateName, DocxTemplateRenderer] = {
    "standard": StandardDocxTemplate(),
    "modern": ModernDocxTemplate(),
    "compact": CompactDocxTemplate(),
}


class ResumeExportService:
    def __init__(self, db: Session) -> None:
        self.resume_version_repository = ResumeVersionRepository(db)
        self.job_description_repository = JobDescriptionRepository(db)

    def export_resume_version_docx(
        self,
        *,
        resume_version_id: int,
        user_id: int,
        template: DocxTemplateName = "standard",
    ) -> BinaryExport:
        resume_version = self.resume_version_repository.get_by_id_for_user(
            resume_version_id=resume_version_id,
            user_id=user_id,
        )
        if resume_version is None:
            raise ExportNotFoundError("Resume version was not found.")

        filename = self._build_export_filename(
            title_source=self._get_export_title_source(resume_version, user_id=user_id),
            resume_version=resume_version,
            template=template,
            extension="docx",
        )
        content = self._build_docx_content(resume_version.content_markdown, template=template)
        return BinaryExport(
            content=content,
            filename=filename,
            content_disposition=self._build_content_disposition(filename),
        )

    def _build_docx_content(self, markdown: str, *, template: DocxTemplateName) -> bytes:
        renderer = DOCX_TEMPLATE_RENDERERS[template]
        return renderer.render(self._parse_markdown(markdown))

    def _parse_markdown(self, markdown: str) -> list[MarkdownBlock]:
        blocks: list[MarkdownBlock] = []
        for line in markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                blocks.append(MarkdownBlock(kind="spacer"))
                continue

            if stripped.startswith("### "):
                blocks.append(MarkdownBlock(kind="heading3", text=stripped.removeprefix("### ").strip()))
                continue
            if stripped.startswith("## "):
                blocks.append(MarkdownBlock(kind="heading2", text=stripped.removeprefix("## ").strip()))
                continue
            if stripped.startswith("# "):
                blocks.append(MarkdownBlock(kind="heading1", text=stripped.removeprefix("# ").strip()))
                continue

            bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
            if bullet_match:
                blocks.append(MarkdownBlock(kind="bullet", text=bullet_match.group(1).strip()))
                continue

            blocks.append(MarkdownBlock(kind="paragraph", text=stripped))
        return blocks

    def _get_export_title_source(self, resume_version: ResumeVersion, *, user_id: int) -> str:
        if resume_version.job_description_id is None:
            return resume_version.title

        job_description = self.job_description_repository.get_by_id_for_user(
            job_description_id=resume_version.job_description_id,
            user_id=user_id,
        )
        if job_description is None:
            return resume_version.title
        return job_description.title

    def _build_export_filename(
        self,
        *,
        title_source: str,
        resume_version: ResumeVersion,
        template: DocxTemplateName,
        extension: str,
    ) -> str:
        date_text = resume_version.created_at.strftime("%Y%m%d")
        safe_title = self._sanitize_filename_part(title_source)
        if not safe_title:
            safe_title = f"ResumeVersion_{resume_version.id}"
        return f"ResumeFit_{safe_title}_{template}_{date_text}.{extension}"

    def _sanitize_filename_part(self, value: str) -> str:
        sanitized = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value)
        sanitized = re.sub(r"\s+", "_", sanitized)
        sanitized = re.sub(r"_+", "_", sanitized).strip(" ._")
        if not sanitized:
            return ""
        if sanitized.upper() in WINDOWS_RESERVED_FILENAMES:
            sanitized = f"{sanitized}_"
        return sanitized[:MAX_EXPORT_FILENAME_STEM_LENGTH].rstrip(" ._")

    def _build_content_disposition(self, filename: str) -> str:
        stem, _, extension = filename.rpartition(".")
        ascii_fallback = stem.encode("ascii", "ignore").decode("ascii")
        ascii_fallback = self._sanitize_filename_part(ascii_fallback) or "ResumeFit"
        quoted_filename = quote(filename)
        return f'attachment; filename="{ascii_fallback}.{extension}"; filename*=UTF-8\'\'{quoted_filename}'
